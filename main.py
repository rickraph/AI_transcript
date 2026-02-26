import os
import io
import shutil
import uuid
from typing import List
import aiofiles
from fastapi import FastAPI, UploadFile, File, Request, Form
from fastapi.responses import HTMLResponse, FileResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from dotenv import load_dotenv
from services import merge_audios, transcribe_audio

# Load .env
load_dotenv()

app = FastAPI(title="Oliveboard Audio Transcriber")

os.makedirs("uploads", exist_ok=True)
os.makedirs("processed", exist_ok=True)
os.makedirs("static", exist_ok=True)

app.mount("/static", StaticFiles(directory="static"), name="static")
templates = Jinja2Templates(directory="templates")


# ─────────────────────────────────────────────
# Helper: extract plain text from docx or pdf
# ─────────────────────────────────────────────
def extract_text(file_bytes: bytes, filename: str) -> str:
    ext = os.path.splitext(filename)[1].lower()

    if ext in (".docx", ".doc"):
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        lines = []
        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text)
        # Also pull text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        lines.append(cell.text.strip())
        return "\n".join(lines)

    elif ext == ".pdf":
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            pages = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
            return "\n".join(pages)
        except ImportError:
            # Fallback: try pypdf
            try:
                from pypdf import PdfReader
                reader = PdfReader(io.BytesIO(file_bytes))
                pages = [p.extract_text() for p in reader.pages if p.extract_text()]
                return "\n".join(pages)
            except ImportError:
                raise RuntimeError(
                    "No PDF library found. Run: pip install PyPDF2  or  pip install pypdf"
                )

    else:
        # Plain text fallback
        return file_bytes.decode("utf-8", errors="ignore")


# ─────────────────────────────────────────────
# Routes
# ─────────────────────────────────────────────
@app.get("/", response_class=HTMLResponse)
async def read_root(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})


@app.post("/process/")
async def process_audio(
    files: List[UploadFile] = File(...),
    file_order: str = Form(...)
):
    try:
        session_id = str(uuid.uuid4())
        session_dir = os.path.join("uploads", session_id)
        os.makedirs(session_dir, exist_ok=True)

        # 1. Save uploaded files
        saved_paths = {}
        for file in files:
            file_path = os.path.join(session_dir, file.filename)
            async with aiofiles.open(file_path, 'wb') as out_file:
                content = await file.read()
                await out_file.write(content)
            saved_paths[file.filename] = file_path

        # 2. Reorder files
        ordered_filenames = [f.strip() for f in file_order.split(',') if f.strip()]
        ordered_paths = [saved_paths[fname] for fname in ordered_filenames if fname in saved_paths]

        # 3. Merge
        merged_file_path = os.path.join("processed", f"{session_id}_merged.mp3")
        await merge_audios(ordered_paths, merged_file_path)

        # 4. Transcribe
        transcription_json, token_usage = await transcribe_audio(merged_file_path)

        # Cleanup
        shutil.rmtree(session_dir)

        return JSONResponse(content={
            "status": "success",
            "merged_audio_url": f"/download/{session_id}_merged.mp3",
            "transcription": transcription_json,
            "token_usage": token_usage
        })

    except Exception as e:
        print(f"SERVER ERROR: {e}")
        return JSONResponse(content={"status": "error", "message": str(e)}, status_code=500)


@app.post("/generate-timeline/")
async def generate_timeline(
    transcription_json: UploadFile = File(...),
    slide_doc: UploadFile = File(...)
):
    import json as _json
    from google import genai
    from google.genai import types

    try:
        session_id = str(uuid.uuid4())
        session_dir = os.path.join("uploads", session_id)
        os.makedirs(session_dir, exist_ok=True)

        # ── Read both files into memory ──
        transcription_bytes = await transcription_json.read()
        slide_doc_bytes = await slide_doc.read()

        # ── Parse transcription JSON ──
        transcription_data = _json.loads(transcription_bytes.decode("utf-8"))
        transcription_text = _json.dumps(transcription_data, indent=2)

        # ── Extract plain text from docx/pdf ──
        # Gemini does NOT support .docx via Files API — we send text instead
        print(f"DEBUG: Extracting text from '{slide_doc.filename}'...")
        slide_text = extract_text(slide_doc_bytes, slide_doc.filename)
        print(f"DEBUG: Extracted {len(slide_text)} characters from slide doc.")

        if not slide_text.strip():
            raise ValueError("Could not extract any text from the slide document. Is it a scanned PDF?")

        # ── Call Gemini — everything as plain text ──
        api_key = os.getenv("GEMINI_API_KEY")
        client = genai.Client(api_key=api_key)

        TIMELINE_PROMPT = open(
            os.path.join(os.path.dirname(__file__), "timeline_prompt.txt")
        ).read()

        full_prompt = (
            f"{TIMELINE_PROMPT}\n\n"
            f"## AUDIO TRANSCRIPTION JSON\n\n```json\n{transcription_text}\n```\n\n"
            f"## SLIDE SCRIPT\n\n{slide_text}"
        )

        print("DEBUG: Sending to Gemini 2.5 Pro...")
        response = client.models.generate_content(
            model="gemini-2.5-pro",
            contents=[
                types.Content(
                    role="user",
                    parts=[types.Part.from_text(text=full_prompt)]
                )
            ],
            config={"response_mime_type": "application/json"}
        )

        # ── Parse result ──
        try:
            result = _json.loads(response.text)
        except _json.JSONDecodeError:
            text = response.text.strip()
            if text.startswith("```"):
                text = "\n".join(text.split("\n")[1:-1])
            result = _json.loads(text)

        # ── Extract token usage ──
        token_usage = None
        try:
            um = response.usage_metadata
            print(f"DEBUG usage_metadata: {um}")
            print(f"DEBUG usage_metadata dir: {dir(um)}")
            # Try all known field name variants
            input_tokens = (
                getattr(um, 'prompt_token_count', None) or
                getattr(um, 'input_tokens', None) or
                getattr(um, 'prompt_tokens', None) or 0
            )
            output_tokens = (
                getattr(um, 'candidates_token_count', None) or
                getattr(um, 'output_tokens', None) or
                getattr(um, 'completion_tokens', None) or 0
            )
            token_usage = {
                "input":  input_tokens,
                "output": output_tokens,
                "model":  "2.5-pro"
            }
            print(f"DEBUG token_usage resolved: {token_usage}")
        except Exception as e:
            print(f"DEBUG token_usage error: {e}")

        # ── Save output ──
        output_path = os.path.join("processed", f"{session_id}_master_plan.json")
        with open(output_path, "w", encoding="utf-8") as f:
            _json.dump(result, f, indent=2, ensure_ascii=False)

        shutil.rmtree(session_dir, ignore_errors=True)

        return JSONResponse(content={
            "status": "success",
            "download_url": f"/download-json/{session_id}_master_plan.json",
            "clip_count": len(result.get("clips", [])),
            "plan": result,
            "token_usage": token_usage
        })

    except Exception as e:
        print(f"TIMELINE ERROR: {e}")
        return JSONResponse(content={"status": "error", "message": str(e)}, status_code=500)


@app.get("/download-json/{filename}")
async def download_json(filename: str):
    file_path = os.path.join("processed", filename)
    if os.path.exists(file_path):
        return FileResponse(file_path, media_type="application/json", filename="master_plan.json")
    return JSONResponse(content={"error": "File not found"}, status_code=404)


@app.get("/download/{filename}")
async def download_file(filename: str):
    file_path = os.path.join("processed", filename)
    if os.path.exists(file_path):
        return FileResponse(file_path, media_type="audio/mpeg", filename="merged_transcript.mp3")
    return JSONResponse(content={"error": "File not found"}, status_code=404)